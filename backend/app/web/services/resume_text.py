"""Plain text out of an uploaded document — ports `server/services/resume.ts`.

PDF, DOCX and plain text. Used by two callers with different needs: the bulk-invite
extractor scans the text for email addresses, and the résumé-driven question
generation feeds it to a model.

Distinct from `app.resume` on the common surface, which asks Gemini to transcribe a
PDF. Both are legitimate — this one is local, free and instant but blind to scanned
images; that one reads scans but costs a model call. The web flow wants the cheap one.

Every extractor returns "" rather than raising: an unreadable or password-protected
file is a normal thing for a recruiter to upload, and the caller's own "no emails
found" message is more useful than a stack trace.
"""

from __future__ import annotations

import asyncio
import io
import logging

logger = logging.getLogger("web.resume_text")

# Bounds what a single upload can cost in memory and parse time.
MAX_BYTES = 10 * 1024 * 1024

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _looks_like(filename: str, *suffixes: str) -> bool:
    lowered = (filename or "").lower()
    return any(lowered.endswith(suffix) for suffix in suffixes)


def _pdf_text(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _docx_text(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [paragraph.text for paragraph in document.paragraphs]
    # Tables too: a CV template often puts contact details in one, and skipping them
    # would silently lose the email address the extractor is looking for.
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _plain_text(data: bytes) -> str:
    # `replace` rather than strict: a CV exported from a Windows tool is often not
    # UTF-8, and one bad byte should not lose the whole document.
    return data.decode("utf-8", "replace")


async def extract(data: bytes, *, content_type: str, filename: str) -> str:
    """Text from a document. Returns "" when nothing can be read.

    Runs in a thread: PDF parsing is CPU-bound and would otherwise block the event
    loop for the length of the document.
    """
    if not data:
        return ""

    mime = (content_type or "").lower()

    if mime == PDF_MIME or _looks_like(filename, ".pdf"):
        extractor = _pdf_text
    elif mime == DOCX_MIME or _looks_like(filename, ".docx"):
        extractor = _docx_text
    elif mime.startswith("text/") or _looks_like(filename, ".txt", ".md"):
        extractor = _plain_text
    else:
        # Unknown type: try plain text rather than refusing. A recruiter's ".dat"
        # export of a contact list is still readable, and the caller validates what
        # comes back.
        extractor = _plain_text

    def _run() -> str:
        try:
            return extractor(data).strip()
        except Exception as exc:  # noqa: BLE001 - an unreadable upload is not an error
            logger.info(
                "could not extract text from %r (%s): %s",
                filename,
                mime,
                type(exc).__name__,
            )
            return ""

    return await asyncio.to_thread(_run)
